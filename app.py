import sys
import os
import json
import logging
from pathlib import Path
from docx import Document
from docx2pdf import convert
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QPalette, QColor
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QFileDialog,
    QLabel,
    QMessageBox,
    QProgressBar,
    QListWidget,
    QStatusBar
)

# Configuración de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Worker(QThread):
    progress = pyqtSignal(int)
    error = pyqtSignal(str)

    def __init__(self, informes, plantilla_path, salida_dir):
        super().__init__()
        self.informes = informes
        self.plantilla_path = Path(plantilla_path)
        self.salida_dir = Path(salida_dir)

    def run(self):
        total = len(self.informes)
        for idx, datos in enumerate(self.informes, start=1):
            nombre_base = f"informe_{datos['estudiante']}_{datos['numero_semana']}"
            nombre_final = self._evitar_choque(nombre_base)
            try:
                datos_pre = self._calcular_hora_total(dict(datos))
                docx_path = self.salida_dir / f"{nombre_final}.docx"
                self._generar_docx(datos_pre, docx_path)

                pdf_path = self.salida_dir / f"{nombre_final}.pdf"
                convert(str(docx_path), str(pdf_path))
                docx_path.unlink(missing_ok=True)
            except Exception as e:
                msg = f"Error en '{nombre_base}': {e}"
                logging.error(msg)
                self.error.emit(msg)

            porcentaje = int(idx / total * 100)
            self.progress.emit(porcentaje)

    def _evitar_choque(self, base_name: str) -> str:
        candidate = base_name
        count = 1
        while (self.salida_dir / f"{candidate}.pdf").exists():
            count += 1
            candidate = f"{base_name}_v{count}"
        return candidate

    def _calcular_hora_total(self, datos: dict) -> dict:
        total = 0.0
        for dia in ["lunes","martes","miercoles","jueves","viernes","sabado"]:
            try:
                total += float(datos.get(f"hora_{dia}", 0))
            except Exception:
                continue
        datos['hora_total'] = str(total)
        return datos

    def _generar_docx(self, datos: dict, destino: Path):
        doc = Document(self.plantilla_path)
        for clave, valor in datos.items():
            marcador = f"[{clave}]"
            for p in doc.paragraphs:
                if marcador in p.text:
                    text = ''.join(run.text for run in p.runs).replace(marcador, str(valor))
                    for run in p.runs:
                        p._element.remove(run._element)
                    p.add_run(text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if marcador in cell.text:
                            cell.text = cell.text.replace(marcador, str(valor))
        doc.save(str(destino))

class InformeGeneratorUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Generador de Informes Semanales")
        self.resize(600, 450)
        self._init_paths()
        self._init_ui()

    def _init_paths(self):
        base = Path.cwd()
        self.dir_json = base / "informe_json"
        self.dir_plantilla = base / "informe_plantilla"
        self.dir_salida = base / "informe_out"
        self.plantilla_docx = self.dir_plantilla / "plantilla.docx"
        self.dir_salida.mkdir(exist_ok=True)

    def _init_ui(self):
        # Central widget y layout
        central = QWidget()
        central.setStyleSheet("background-color: #2b2b2b; color: #ffffff;")
        layout = QVBoxLayout(central)
        layout.setContentsMargins(12,12,12,12)
        layout.setSpacing(10)

        # Botones superiores
        botones = QHBoxLayout()
        self.btn_cargar = QPushButton("📂 Cargar JSON")
        self.btn_cargar.setCursor(Qt.PointingHandCursor)
        self.btn_cargar.clicked.connect(self.cargar_json)
        self.btn_generar = QPushButton("🚀 Generar")
        self.btn_generar.setCursor(Qt.PointingHandCursor)
        self.btn_generar.setEnabled(False)
        self.btn_generar.clicked.connect(self.generar_informes)
        botones.addWidget(self.btn_cargar)
        botones.addWidget(self.btn_generar)
        layout.addLayout(botones)

        # Lista de informes
        layout.addWidget(QLabel("Informes a generar:"))
        self.lista = QListWidget()
        self.lista.setStyleSheet(
            "background-color: #313335; border: 1px solid #5A5A5A; border-radius: 5px;"
        )
        layout.addWidget(self.lista)

        # Barra de progreso global
        layout.addWidget(QLabel("Progreso global:"))
        self.barra = QProgressBar()
        self.barra.setAlignment(Qt.AlignCenter)
        self.barra.setStyleSheet(
            "QProgressBar { background-color: #3c3f41; border: 1px solid #5A5A5A; border-radius: 5px; text-align: center;}"
            "QProgressBar::chunk { background-color: #61afef; border-radius: 5px;}"
        )
        layout.addWidget(self.barra)

        self.setCentralWidget(central)

        # Status bar
        self.status = QStatusBar()
        self.status.setStyleSheet("background-color: #212121; color: #aaaaaa;")
        self.setStatusBar(self.status)

    def cargar_json(self):
        ruta, _ = QFileDialog.getOpenFileName(
            self, "Seleccionar JSON", str(self.dir_json), "JSON Files (*.json)"
        )
        if not ruta:
            return
        try:
            with open(ruta, encoding='utf-8') as f:
                self.informes = json.load(f)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo leer el JSON:\n{e}")
            return

        self.lista.clear()
        for datos in self.informes:
            nombre = f"informe_{datos['estudiante']}_{datos['numero_semana']}"
            self.lista.addItem(nombre)
        self.btn_generar.setEnabled(True)
        self.status.showMessage(f"{len(self.informes)} informes listos para generar.")
        self.barra.setValue(0)

    def generar_informes(self):
        self.btn_generar.setEnabled(False)
        self.worker = Worker(self.informes, self.plantilla_docx, self.dir_salida)
        self.worker.progress.connect(self.barra.setValue)
        self.worker.error.connect(lambda msg: QMessageBox.warning(self, "Error en generación", msg))
        self.worker.finished.connect(lambda: QMessageBox.information(self, "¡Completado!", "Todos los informes se han generado en PDF."))
        self.worker.finished.connect(lambda: self.btn_generar.setEnabled(True))
        self.worker.start()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    dark_palette = QPalette()
    dark_palette.setColor(QPalette.Window, QColor(43,43,43))
    dark_palette.setColor(QPalette.WindowText, QColor(220,220,220))
    dark_palette.setColor(QPalette.Base, QColor(35,35,35))
    dark_palette.setColor(QPalette.AlternateBase, QColor(43,43,43))
    dark_palette.setColor(QPalette.ToolTipBase, QColor(255,255,220))
    dark_palette.setColor(QPalette.ToolTipText, QColor(220,220,220))
    dark_palette.setColor(QPalette.Text, QColor(220,220,220))
    dark_palette.setColor(QPalette.Button, QColor(43,43,43))
    dark_palette.setColor(QPalette.ButtonText, QColor(220,220,220))
    dark_palette.setColor(QPalette.BrightText, Qt.red)
    dark_palette.setColor(QPalette.Highlight, QColor(61,174,239))
    dark_palette.setColor(QPalette.HighlightedText, QColor(35,35,35))
    app.setPalette(dark_palette)
    ventana = InformeGeneratorUI()
    ventana.show()
    sys.exit(app.exec_())
