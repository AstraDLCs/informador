#!/usr/bin/env python3
"""
TUI para generación de informes semanales con Textual y Rich,
compatible con Termux/Linux. Solo guarda los informes en .docx.
"""
import os
import json
import asyncio
from pathlib import Path
from rich.console import Console
from textual.app import App, ComposeResult
from textual.widgets import Header, Footer, Button, ListView, ListItem, Static, ProgressBar, Label
from textual.containers import Horizontal

from docx import Document

# ─── RUTAS ────────────────────────────────────────────────────────────────────────────
BASE       = Path.cwd()
DIR_JSON   = BASE / "informe_json";     DIR_JSON.mkdir(exist_ok=True)
PLANTILLA  = BASE / "informe_plantilla" / "plantilla.docx"
DIR_OUT    = BASE / "informe_out";       DIR_OUT.mkdir(exist_ok=True)

console = Console()

def sanitizar_id(nombre: str) -> str:
    """
    Reemplaza '.' por '_' y antepone '_' si comienza con dígito,
    para cumplir ^[A-Za-z_][A-Za-z0-9_-]*$ en Textual.
    """
    valid = nombre.replace(".", "_")
    if valid and valid[0].isdigit():
        valid = f"_{valid}"
    return valid

class InformeApp(App):
    CSS = """
    Screen { background: #1e1e1e; color: white; }
    ListView { height: 8; border: round yellow; }
    ProgressBar { height: 1; width: 80;
      border: round #5A5A5A; background: #3c3f41; }
    Button { margin: 1 2; background: #2b2b2b; color: white; }
    """

    def compose(self) -> ComposeResult:
        yield Header()
        yield Static("[b]Selecciona un archivo JSON:[/b]")

        # Listar JSON y crear ListItem(Label, id=sanitizado)
        files = [f for f in os.listdir(DIR_JSON) if f.endswith(".json")]
        items = [ListItem(Label(f), id=sanitizar_id(f)) for f in files]
        yield ListView(*items, id="json_list")

        with Horizontal():
            yield Button("Cambiar JSON", id="btn_change")
            yield Button("Generar Informes", id="btn_generate", disabled=True)

        yield Static("[b]Informes a generar:[/b]")
        yield ListView(id="reports_list")

        yield ProgressBar(total=100, id="progress")
        yield Footer()

    async def on_list_view_selected(self, event: ListView.Selected) -> None:
        # Habilitar botón "Generar Informes"
        self.query_one("#btn_generate", Button).disabled = False

        # Recuperar nombre del fichero JSON según event.item.id
        sel_id = event.item.id
        for fname in os.listdir(DIR_JSON):
            if sanitizar_id(fname) == sel_id:
                path = DIR_JSON / fname
                break

        # Cargar datos del JSON
        with open(path, encoding="utf-8") as f:
            self.current_data = json.load(f)

        # Mostrar lista de informes a crear
        reports_list = self.query_one("#reports_list", ListView)
        reports_list.clear()
        for datos in self.current_data:
            name = f"informe_{datos['estudiante']}_{datos['numero_semana']}"
            reports_list.append(ListItem(Label(name)))

    async def on_button_pressed(self, event: Button.Pressed) -> None:
        btn = event.button.id
        if btn == "btn_change":
            # Reset interfaz
            self.query_one("#json_list", ListView).index = None
            self.query_one("#reports_list", ListView).clear()
            self.query_one("#btn_generate", Button).disabled = True
            self.query_one("#progress", ProgressBar).update(progress=0)
        elif btn == "btn_generate":
            await self.generate_reports()

    async def generate_reports(self):
        progress = self.query_one("#progress", ProgressBar)
        total = len(self.current_data)

        for idx, datos in enumerate(self.current_data, start=1):
            # Nombre base y evitar choques (_v2, _v3...)
            base = f"informe_{datos['estudiante']}_{datos['numero_semana']}"
            nombre = self._no_conflicto(base)

            # Calcular hora_total
            horas = sum(
                float(datos.get(f"hora_{d}", 0))
                for d in ["lunes", "martes", "miercoles", "jueves", "viernes", "sabado"]
            )
            datos["hora_total"] = str(horas)

            # Generar y guardar DOCX
            salida_docx = DIR_OUT / f"{nombre}.docx"
            doc = Document(str(PLANTILLA))
            for k, v in datos.items():
                ph = f"[{k}]"
                # Reemplazo en párrafos
                for p in doc.paragraphs:
                    if ph in p.text:
                        text = "".join(run.text for run in p.runs).replace(ph, str(v))
                        for run in p.runs:
                            p._element.remove(run._element)
                        p.add_run(text)
                # Reemplazo en tablas
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            if ph in cell.text:
                                cell.text = cell.text.replace(ph, str(v))
            doc.save(str(salida_docx))

            # Actualizar barra de progreso
            progreso = int(idx / total * 100)
            progress.update(progress=progreso)
            await asyncio.sleep(0.05)

        console.log("[green]Generación de .docx completada.[/green]")
        self.exit()

    def _no_conflicto(self, base: str) -> str:
        """Evita sobreescritura añadiendo sufijos _v2, _v3…"""
        i = 1
        nombre = base
        while (DIR_OUT / f"{nombre}.docx").exists():
            i += 1
            nombre = f"{base}_v{i}"
        return nombre

if __name__ == "__main__":
    InformeApp().run()
